import pandas as pd
from sklearn.svm import SVC
from sklearn.feature_extraction.text import TfidfVectorizer
from docx import Document

# ====== CONFIG ======
TRAINING_EXCEL = r"C:\Users\adij2\Downloads\clean_full_training_dataset.xlsx"
INPUT_DOCX = r"C:\Users\adij2\Downloads\sample - Input.docx"
OUTPUT_DOCX = r"C:\Users\adij2\Downloads\ProcessedLabels.docx"
GT_DOCX = "C:\\Users\\adij2\\Downloads\\GroundTruth.docx"

# ====== Load training data ======
data = pd.read_excel(TRAINING_EXCEL)
X_train_text = (data['raw_label'].astype(str) + " " + data['description'].astype(str))
y_train = data['canonical_label'].astype(str)

# ====== Train TF-IDF vectorizer + SVM ======
vectorizer = TfidfVectorizer(analyzer='char_wb', ngram_range=(2, 4))
X_train_tfidf = vectorizer.fit_transform(X_train_text)

model = SVC(kernel='rbf', C=1.0, gamma='scale', probability=True)
model.fit(X_train_tfidf, y_train)

# ====== Load input Word doc ======
raw_doc = Document(INPUT_DOCX)
all_lines = [p.text.rstrip() for p in raw_doc.paragraphs if p.text.strip() != ""]

# ====== Process lines: alternate raw_label and description ======
output_lines = []
i = 0
while i < len(all_lines):
    raw_label_line = all_lines[i].strip()
    description_line = all_lines[i+1].strip() if (i+1) < len(all_lines) else ""

    # Combine raw_label + description for prediction
    combined_input = f"{raw_label_line} {description_line}".strip()
    vec = vectorizer.transform([combined_input])
    pred_label = model.predict(vec)[0]

    # Append canonical label + original description
    output_lines.append(pred_label)
    output_lines.append(description_line)

    i += 2  # move to next raw_label/description pair

# ====== Save processed doc ======
processed_doc = Document()
for ln in output_lines:
    processed_doc.add_paragraph(ln)
processed_doc.save(OUTPUT_DOCX)
print("Saved processed file:", OUTPUT_DOCX)

raw_labels_set = set(raw_df["raw_label"].astype(str))

# Load existing GroundTruth doc
gt_doc = Document(GT_DOCX)

# If a table exists, use it; otherwise create a new table with headers
if gt_doc.tables:
    table = gt_doc.tables[0]
else:
    table = gt_doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Ground Truth"
    table.rows[0].cells[1].text = "Found/Not Found"

# Collect existing ground truth labels in the table
existing_labels = set()
for row in table.rows[1:]:  # skip header
    gt_label = row.cells[0].text.strip()
    existing_labels.add(gt_label)

# Gather all labels to process: existing + any new ones from elsewhere
# Here, just re-using existing table cells (or could append new ones)
for row in table.rows[1:]:
    gt_label = row.cells[0].text.strip()
    row.cells[1].text = "Found" if gt_label in raw_labels_set else "Not Found"

gt_doc.save(GT_DOCX)
print("Ground Truth updated with Found/Not Found")
