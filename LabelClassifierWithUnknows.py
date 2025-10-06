import pandas as pd
import numpy as np
from sklearn.svm import SVC
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.model_selection import train_test_split
from docx import Document

# ====== CONFIG ======
TRAINING_EXCEL = r"C:\Users\adij2\Downloads\clean_full_training_dataset.xlsx"
INPUT_DOCX = r"C:\Users\adij2\Downloads\sample - Input.docx"
OUTPUT_DOCX = r"C:\Users\adij2\Downloads\ProcessedLabels.docx"
GT_DOCX = "C:\\Users\\adij2\\Downloads\\GroundTruth.docx"
UNKNOWN_PERCENTILE = 20  # bottom X% confidence will be UNKNOWN_LABEL

# ====== Load training data ======
data = pd.read_excel(TRAINING_EXCEL)
X_text = (data['raw_label'].astype(str) + " " + data['description'].astype(str))
y = data['canonical_label'].astype(str)

# ====== Split for threshold calibration ======
X_train_text, X_val_text, y_train, y_val = train_test_split(
    X_text, y, test_size=0.2, stratify=y, random_state=42
)

# ====== Train TF-IDF vectorizer + SVM ======
vectorizer = TfidfVectorizer(analyzer='char_wb', ngram_range=(2, 4))
X_train_tfidf = vectorizer.fit_transform(X_train_text)
X_val_tfidf = vectorizer.transform(X_val_text)

model = SVC(kernel='rbf', C=1.0, gamma='scale', probability=True)
model.fit(X_train_tfidf, y_train)

# ====== Compute validation max probabilities ======
val_probs = model.predict_proba(X_val_tfidf)
val_max = val_probs.max(axis=1)

# ====== Set percentile-based threshold ======
PROB_THRESHOLD = np.percentile(val_max, UNKNOWN_PERCENTILE)
print(f"Using probability threshold = {PROB_THRESHOLD:.2f}")

# ====== Load input Word doc ======
raw_doc = Document(INPUT_DOCX)
all_lines = [p.text.rstrip() for p in raw_doc.paragraphs if p.text.strip() != ""]

# ====== Process lines: alternate raw_label and description ======
output_lines = []
i = 0
while i < len(all_lines):
    raw_label_line = all_lines[i].strip()
    description_line = all_lines[i+1].strip() if (i+1) < len(all_lines) else ""

    # Combine raw_label + description for prediction
    combined_input = f"{raw_label_line} {description_line}".strip()
    vec = vectorizer.transform([combined_input])
    probs = model.predict_proba(vec)[0]
    max_prob = probs.max()
    pred_label = model.classes_[probs.argmax()]

    # Apply threshold: mark as UNKNOWN if low confidence
    if max_prob < PROB_THRESHOLD:
        pred_label = "UNKNOWN_LABEL"

    # Append predicted label + original description
    output_lines.append(pred_label)
    output_lines.append(description_line)

    i += 2  # move to next raw_label/description pair

# ====== Save processed doc ======
processed_doc = Document()
for ln in output_lines:
    processed_doc.add_paragraph(ln)
processed_doc.save(OUTPUT_DOCX)
print("Saved processed file:", OUTPUT_DOCX)

raw_labels_set = set(raw_df["raw_label"].astype(str))

# Load existing GroundTruth doc
gt_doc = Document(GT_DOCX)

# If a table exists, use it; otherwise create a new table with headers
if gt_doc.tables:
    table = gt_doc.tables[0]
else:
    table = gt_doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Ground Truth"
    table.rows[0].cells[1].text = "Found/Not Found"

# Collect existing ground truth labels in the table
existing_labels = set()
for row in table.rows[1:]:  # skip header
    gt_label = row.cells[0].text.strip()
    existing_labels.add(gt_label)

# Gather all labels to process: existing + any new ones from elsewhere
# Here, just re-using existing table cells (or could append new ones)
for row in table.rows[1:]:
    gt_label = row.cells[0].text.strip()
    row.cells[1].text = "Found" if gt_label in raw_labels_set else "Not Found"

gt_doc.save(GT_DOCX)
print("Ground Truth updated with Found/Not Found")
